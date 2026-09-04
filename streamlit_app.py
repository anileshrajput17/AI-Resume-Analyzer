import streamlit as st
import io
import re

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="🤖",
    layout="centered"
)

# =========================
# TEXT EXTRACTION
# =========================

def extract_text(uploaded_file):
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    # ---------- PDF ----------
    if file_name.endswith(".pdf"):
        try:
            import fitz

            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = []

            for page in doc:
                page_text = page.get_text("text").strip()

                # Normal text PDF
                if page_text:
                    pages_text.append(page_text)
                    continue

                # Scanned/image PDF OCR
                try:
                    import pytesseract
                    from PIL import Image

                    pix = page.get_pixmap(
                        matrix=fitz.Matrix(2, 2),
                        alpha=False
                    )

                    img = Image.open(
                        io.BytesIO(pix.tobytes("png"))
                    )

                    ocr_text = pytesseract.image_to_string(img)

                    if ocr_text.strip():
                        pages_text.append(ocr_text)

                except Exception:
                    pass

            doc.close()

            return "\n".join(pages_text)

        except Exception:
            return ""

    # ---------- DOCX ----------
    if file_name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))

            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Also read tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            return "\n".join(text_parts)

        except Exception:
            return ""

    return ""


# =========================
# ANALYSIS
# =========================

def analyze_resume(text):

    text_lower = text.lower()

    # ---------- SKILLS ----------
    skills = {
        "Python": r"\bpython\b",
        "Java": r"\bjava\b",
        "JavaScript": r"\bjavascript\b",
        "HTML": r"\bhtml\b",
        "CSS": r"\bcss\b",
        "SQL": r"\bsql\b",
        "Excel": r"\bexcel\b",
        "PowerPoint": r"\bpowerpoint\b",
        "Word": r"\bms word\b|\bword\b",
        "Communication": r"\bcommunication\b",
        "Leadership": r"\bleadership\b",
        "Teamwork": r"\bteamwork\b|\bteam work\b",
        "Problem Solving": r"\bproblem solving\b",
        "Data Analysis": r"\bdata analysis\b",
        "Machine Learning": r"\bmachine learning\b",
        "Artificial Intelligence": r"\bartificial intelligence\b|\bai\b",
        "Project Management": r"\bproject management\b",
        "Marketing": r"\bmarketing\b",
        "Research": r"\bresearch\b",
        "Management": r"\bmanagement\b",
        "C++": r"\bc\+\+\b",
        "C": r"\bc programming\b|\blanguage c\b",
        "Git": r"\bgit\b",
        "GitHub": r"\bgithub\b",
        "Power BI": r"\bpower bi\b",
        "Communication Skills": r"\bcommunication skills\b"
    }

    found_skills = []

    for skill, pattern in skills.items():
        if re.search(pattern, text_lower):
            found_skills.append(skill)

    # ---------- SECTIONS ----------

    section_patterns = {

        "education": [
            r"\beducation\b",
            r"\bacademic qualification\b",
            r"\bqualification\b",
            r"\bdegree\b",
            r"\bbachelor\b",
            r"\bmaster\b",
            r"\bcollege\b",
            r"\buniversity\b",
            r"\bschool\b"
        ],

        "experience": [
            r"\bexperience\b",
            r"\bwork experience\b",
            r"\bprofessional experience\b",
            r"\bemployment\b",
            r"\binternship\b",
            r"\bworked\b"
        ],

        "projects": [
            r"\bprojects\b",
            r"\bproject\b",
            r"\bacademic projects\b",
            r"\bpersonal projects\b"
        ],

        "certifications": [
            r"\bcertification\b",
            r"\bcertifications\b",
            r"\bcertificate\b",
            r"\bcertificates\b"
        ],

        "summary": [
            r"\bsummary\b",
            r"\bprofessional summary\b",
            r"\bcareer objective\b",
            r"\bobjective\b",
            r"\bprofile\b",
            r"\babout me\b"
        ],

        "contact": [
            r"\bemail\b",
            r"\bphone\b",
            r"\bmobile\b",
            r"\blinkedin\b",
            r"\bgithub\b"
        ]
    }

    sections = {}

    for section, patterns in section_patterns.items():
        sections[section] = any(
            re.search(pattern, text_lower)
            for pattern in patterns
        )

    # =========================
    # DYNAMIC SCORING
    # =========================

    score = 0

    # Education: 15
    if sections["education"]:
        score += 15

    # Skills: 25
    skill_score = min(len(found_skills) * 2.5, 25)
    score += skill_score

    # Experience: 20
    if sections["experience"]:
        score += 20

    # Projects: 15
    if sections["projects"]:
        score += 15

    # Certifications: 10
    if sections["certifications"]:
        score += 10

    # Summary: 5
    if sections["summary"]:
        score += 5

    # Contact details: 5
    if sections["contact"]:
        score += 5

    score = min(round(score), 100)

    # =========================
    # SUGGESTIONS
    # =========================

    suggestions = []

    if not sections["summary"]:
        suggestions.append(
            "Add a professional summary or career objective."
        )

    if not sections["education"]:
        suggestions.append(
            "Add a clear Education section."
        )

    if len(found_skills) < 5:
        suggestions.append(
            "Add more relevant technical and soft skills."
        )

    if not sections["experience"]:
        suggestions.append(
            "Add work experience, internship or relevant practical experience."
        )

    if not sections["projects"]:
        suggestions.append(
            "Add academic or personal projects."
        )

    if not sections["certifications"]:
        suggestions.append(
            "Add relevant certifications if you have them."
        )

    if not sections["contact"]:
        suggestions.append(
            "Make sure your contact information is clearly visible."
        )

    return {
        "score": score,
        "skills": found_skills,
        "sections": sections,
        "suggestions": suggestions
    }


# =========================
# USER INTERFACE
# =========================

st.title("🤖 AI Resume Analyzer")

st.write(
    "Upload your resume and get an automated analysis "
    "based on skills, education, experience, projects, "
    "certifications and resume structure."
)

st.divider()

resume = st.file_uploader(
    "📄 Upload your Resume",
    type=["pdf", "docx"]
)

if resume is not None:

    st.success("Resume uploaded successfully! ✅")

    with st.spinner("🔍 Reading and analyzing your resume..."):
        text = extract_text(resume)

    if not text.strip():

        st.error(
            "❌ I couldn't extract readable text from this file."
        )

        st.info(
            "Try a text-based PDF/DOCX. Scanned PDFs require OCR support."
        )

    else:

        result = analyze_resume(text)

        # ---------- SCORE ----------
        st.subheader("📊 Resume Analysis")

        st.metric(
            "Resume Score",
            f"{result['score']}/100"
        )

        st.divider()

        # ---------- SECTIONS ----------
        st.write("### 📋 Resume Sections")

        section_names = {
            "education": "Education",
            "experience": "Experience",
            "projects": "Projects",
            "certifications": "Certifications",
            "summary": "Professional Summary",
            "contact": "Contact Information"
        }

        for key, name in section_names.items():

            if result["sections"][key]:
                st.write(f"✅ {name} detected")
            else:
                st.write(f"⚠️ {name} not detected")

        # ---------- SKILLS ----------
        st.write("### 🛠️ Skills Detected")

        if result["skills"]:
            st.write(", ".join(result["skills"]))
        else:
            st.write("No common skills detected.")

        # ---------- TEXT INFO ----------
        st.write("### 📄 Resume Information")

        word_count = len(text.split())
        character_count = len(text)

        col1, col2 = st.columns(2)

        with col1:
            st.metric("Words", word_count)

        with col2:
            st.metric("Characters", character_count)

        # ---------- SUGGESTIONS ----------
        st.write("### 💡 Improvement Suggestions")

        if result["suggestions"]:

            for suggestion in result["suggestions"]:
                st.write("• " + suggestion)

        else:
            st.success(
                "🎉 Your resume contains the major sections!"
            )

        # ---------- EXTRACTED TEXT ----------
        with st.expander("🔎 View Extracted Resume Text"):
            st.text(text[:10000])

else:

    st.info(
        "Please upload a PDF or DOCX resume to begin."
    )
